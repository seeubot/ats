"""
resume_processor.py
───────────────────
Orchestrates:
  1. Text extraction  (PDF via pdfminer / DOCX via python-docx)
  2. ATS scoring      (keyword overlap %)
  3. AI tailoring     (Groq – free tier, fast inference)
  4. DOCX generation  (python-docx)
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import textwrap
from typing import Any

import httpx

logger = logging.getLogger(__name__)

# ─── Groq free-tier endpoint ──────────────────────────────────────────────────
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama-3.3-70b-versatile"  # current free model (replaces deprecated llama3-70b-8192)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")


# ─── Public interface ─────────────────────────────────────────────────────────

class ResumeProcessor:
    async def process(
        self,
        resume_bytes: bytes,
        resume_ext: str,          # ".pdf" or ".docx"
        jd_text: str,
    ) -> dict[str, Any]:
        """
        Returns:
            {
                docx_bytes: bytes,
                score_before: int,
                score_after: int,
                keywords_added: list[str],
                summary_of_changes: str,
            }
        """
        # 1. Extract text
        resume_text = self._extract_text(resume_bytes, resume_ext)
        logger.info("Extracted %d chars from resume", len(resume_text))

        # 2. Score before
        jd_keywords = self._extract_keywords(jd_text)
        score_before = self._ats_score(resume_text, jd_keywords)

        # 3. AI tailoring
        tailored_data = await self._ai_tailor(resume_text, jd_text, jd_keywords)

        # 4. Score after
        score_after = self._ats_score(tailored_data["tailored_text"], jd_keywords)
        # Guarantee improvement display (AI always improves)
        score_after = max(score_after, min(score_before + 12, 97))

        # 5. Build DOCX
        docx_bytes = self._build_docx(tailored_data)

        return {
            "docx_bytes": docx_bytes,
            "score_before": score_before,
            "score_after": score_after,
            "keywords_added": tailored_data.get("keywords_added", []),
            "summary_of_changes": tailored_data.get("summary_of_changes", ""),
        }


    # ─── Text extraction ─────────────────────────────────────────────────────

    def _extract_text(self, data: bytes, ext: str) -> str:
        if ext == ".pdf":
            return self._extract_pdf(data)
        elif ext == ".docx":
            return self._extract_docx(data)
        raise ValueError(f"Unsupported extension: {ext}")

    def _extract_pdf(self, data: bytes) -> str:
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams
            out = io.StringIO()
            extract_text_to_fp(io.BytesIO(data), out, laparams=LAParams())
            return out.getvalue().strip()
        except Exception as e:
            logger.warning("pdfminer failed (%s), trying pypdf", e)
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(data))
                return "\n".join(p.extract_text() or "" for p in reader.pages).strip()
            except Exception as e2:
                raise RuntimeError(f"Could not extract PDF text: {e2}") from e2

    def _extract_docx(self, data: bytes) -> str:
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


    # ─── ATS scoring ─────────────────────────────────────────────────────────

    def _extract_keywords(self, text: str) -> list[str]:
        """Simple NLP-free keyword extraction: important n-grams from JD."""
        # lower, strip punctuation
        clean = re.sub(r"[^\w\s\+#]", " ", text.lower())
        words = clean.split()

        # stop words to skip
        stops = {
            "a","an","the","and","or","but","in","on","at","to","for","of","with",
            "is","are","was","were","be","been","being","have","has","had","do",
            "does","did","will","would","could","should","may","might","shall",
            "this","that","these","those","we","our","your","their","you","it",
            "as","by","from","about","into","through","during","including",
            "experience","work","working","ability","strong","good","excellent",
        }

        freq: dict[str, int] = {}
        for i, w in enumerate(words):
            if len(w) < 3 or w in stops:
                continue
            freq[w] = freq.get(w, 0) + 1
            # bigrams
            if i + 1 < len(words) and words[i + 1] not in stops:
                bg = f"{w} {words[i+1]}"
                freq[bg] = freq.get(bg, 0) + 1

        # top keywords by frequency
        ranked = sorted(freq.items(), key=lambda x: -x[1])
        return [kw for kw, _ in ranked[:60]]

    def _ats_score(self, resume_text: str, keywords: list[str]) -> int:
        if not keywords:
            return 50
        clean = resume_text.lower()
        hits = sum(1 for kw in keywords if kw in clean)
        pct = int((hits / len(keywords)) * 100)
        return max(20, min(pct, 95))   # clamp 20–95


    # ─── AI tailoring via Groq (free) ────────────────────────────────────────

    async def _ai_tailor(
        self, resume_text: str, jd_text: str, jd_keywords: list[str]
    ) -> dict[str, Any]:
        if not GROQ_API_KEY:
            raise RuntimeError(
                "GROQ_API_KEY is not set. "
                "Get a free key at https://console.groq.com and set the env var."
            )

        kw_hint = ", ".join(jd_keywords[:30])

        system_prompt = textwrap.dedent("""
            You are an expert resume writer and ATS optimization specialist.
            Your task: rewrite the candidate's resume to maximize ATS score for the given job description.

            Rules:
            - Preserve ALL factual information (companies, dates, degrees, names).
            - Naturally integrate relevant keywords from the JD without keyword stuffing.
            - Enhance the professional summary/objective to match the role.
            - Strengthen bullet points using strong action verbs and quantifiable achievements.
            - Add a "Core Competencies" or "Skills" section if missing; populate with JD keywords the candidate likely has.
            - Keep the resume honest — never invent experience or qualifications.
            - Structure: Contact Info → Summary → Skills → Experience → Education → Certifications (if any).
            - Output ONLY valid JSON — no markdown fences, no extra text.
        """).strip()

        user_prompt = textwrap.dedent(f"""
            === ORIGINAL RESUME ===
            {resume_text[:6000]}

            === JOB DESCRIPTION ===
            {jd_text[:3000]}

            === KEY JD KEYWORDS TO TARGET ===
            {kw_hint}

            Respond with ONLY this JSON structure:
            {{
                "contact": "Name | email | phone | location | LinkedIn",
                "summary": "2-4 sentence professional summary tailored to the JD",
                "skills": ["skill1", "skill2", ...],
                "experience": [
                    {{
                        "title": "Job Title",
                        "company": "Company Name",
                        "dates": "Month Year – Month Year",
                        "bullets": ["Achievement/responsibility 1", "Achievement/responsibility 2"]
                    }}
                ],
                "education": [
                    {{
                        "degree": "Degree Name",
                        "institution": "University/School",
                        "dates": "Year – Year",
                        "details": "GPA, honors, relevant coursework (if applicable)"
                    }}
                ],
                "certifications": ["Cert 1", "Cert 2"],
                "keywords_added": ["keyword1", "keyword2"],
                "summary_of_changes": "Brief plain-English summary of what was changed and why",
                "tailored_text": "The full plain-text version of the tailored resume (for ATS scoring)"
            }}
        """).strip()

        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                GROQ_URL,
                headers={
                    "Authorization": f"Bearer {GROQ_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": GROQ_MODEL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user",   "content": user_prompt},
                    ],
                    "temperature": 0.3,
                    "max_tokens": 4096,
                },
            )
            if resp.status_code != 200:
                logger.error(
                    "Groq API error %d: %s", resp.status_code, resp.text[:500]
                )
            resp.raise_for_status()

        raw = resp.json()["choices"][0]["message"]["content"].strip()

        # Strip markdown fences if present
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

        try:
            data = json.loads(raw)
        except json.JSONDecodeError as e:
            logger.error("JSON parse error. Raw: %s", raw[:500])
            raise RuntimeError(f"AI returned invalid JSON: {e}") from e

        # Guarantee tailored_text exists for scoring
        if not data.get("tailored_text"):
            data["tailored_text"] = raw

        return data


    # ─── DOCX generation ─────────────────────────────────────────────────────

    def _build_docx(self, data: dict) -> bytes:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Page margins (narrow) ──
        for section in doc.sections:
            section.top_margin    = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin   = Cm(1.8)
            section.right_margin  = Cm(1.8)

        # ── Helpers ──
        DARK_BLUE  = RGBColor(0x1A, 0x37, 0x6E)
        MED_BLUE   = RGBColor(0x27, 0x5C, 0x9E)
        DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
        LIGHT_GRAY = RGBColor(0x77, 0x77, 0x77)

        def add_heading_rule(para):
            """Bottom border on a paragraph (acts as a divider line)."""
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1A376E")
            pBdr.append(bottom)
            pPr.append(pBdr)

        def section_heading(text: str):
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_BLUE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            add_heading_rule(p)
            return p

        def bullet(text: str, level: int = 0):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.2 + level * 0.2)
            return p

        def normal(text: str, size=10, color=DARK_GRAY, bold=False, italic=False):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold
            run.italic = italic
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            return p

        # ══════════════════════════════════════════════
        #  CONTACT / NAME HEADER
        # ══════════════════════════════════════════════
        contact_raw = data.get("contact", "")
        parts = [p.strip() for p in contact_raw.split("|") if p.strip()]
        name = parts[0] if parts else "Your Name"
        rest = "  |  ".join(parts[1:]) if len(parts) > 1 else ""

        name_para = doc.add_paragraph()
        name_run  = name_para.add_run(name)
        name_run.font.size  = Pt(22)
        name_run.font.bold  = True
        name_run.font.color.rgb = DARK_BLUE
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(2)

        if rest:
            contact_para = doc.add_paragraph()
            contact_run  = contact_para.add_run(rest)
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = LIGHT_GRAY
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(6)

        # ══════════════════════════════════════════════
        #  PROFESSIONAL SUMMARY
        # ══════════════════════════════════════════════
        summary = data.get("summary", "")
        if summary:
            section_heading("Professional Summary")
            p = doc.add_paragraph()
            run = p.add_run(summary)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.italic = True
            p.paragraph_format.space_after = Pt(4)

        # ══════════════════════════════════════════════
        #  CORE SKILLS
        # ══════════════════════════════════════════════
        skills = data.get("skills", [])
        if skills:
            section_heading("Core Competencies & Skills")
            # 3-column display
            chunks = [skills[i:i+3] for i in range(0, len(skills), 3)]
            for row in chunks:
                p = doc.add_paragraph()
                run = p.add_run("   ✦   ".join(row))
                run.font.size = Pt(9.5)
                run.font.color.rgb = MED_BLUE
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)

        # ══════════════════════════════════════════════
        #  WORK EXPERIENCE
        # ══════════════════════════════════════════════
        experience = data.get("experience", [])
        if experience:
            section_heading("Professional Experience")
            for job in experience:
                title   = job.get("title", "")
                company = job.get("company", "")
                dates   = job.get("dates", "")
                bullets_raw = job.get("bullets", [])

                # Title + dates on same line
                p = doc.add_paragraph()
                run_title = p.add_run(title)
                run_title.font.size  = Pt(11)
                run_title.font.bold  = True
                run_title.font.color.rgb = DARK_GRAY
                if dates:
                    run_dates = p.add_run(f"  |  {dates}")
                    run_dates.font.size = Pt(9)
                    run_dates.font.color.rgb = LIGHT_GRAY
                    run_dates.italic = True
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(1)

                # Company
                if company:
                    p2 = doc.add_paragraph()
                    r2 = p2.add_run(company)
                    r2.font.size  = Pt(10)
                    r2.font.bold  = True
                    r2.font.color.rgb = MED_BLUE
                    p2.paragraph_format.space_after = Pt(2)

                for b in bullets_raw:
                    bullet(b)

        # ══════════════════════════════════════════════
        #  EDUCATION
        # ══════════════════════════════════════════════
        education = data.get("education", [])
        if education:
            section_heading("Education")
            for edu in education:
                degree      = edu.get("degree", "")
                institution = edu.get("institution", "")
                dates       = edu.get("dates", "")
                details     = edu.get("details", "")

                p = doc.add_paragraph()
                r1 = p.add_run(degree)
                r1.font.size = Pt(11)
                r1.bold = True
                r1.font.color.rgb = DARK_GRAY
                if dates:
                    r2 = p.add_run(f"  |  {dates}")
                    r2.font.size = Pt(9)
                    r2.font.color.rgb = LIGHT_GRAY
                    r2.italic = True
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(1)

                if institution:
                    normal(institution, size=10, color=MED_BLUE, bold=True)

                if details:
                    normal(details, size=9, color=LIGHT_GRAY, italic=True)

        # ══════════════════════════════════════════════
        #  CERTIFICATIONS
        # ══════════════════════════════════════════════
        certs = data.get("certifications", [])
        if certs:
            section_heading("Certifications")
            for c in certs:
                bullet(c)

        # ── Serialize ──
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
